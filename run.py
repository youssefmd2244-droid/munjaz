#!/usr/bin/env python3
"""
╔══════════════════════════════════════════════════════════════════════════════╗
║   ICON SWARM ULTIMATE v10.0 — Full Video Generation & Storage             ║
║   Multi‑Model AI • Self‑Evolving • Multi‑Storage • VideoSwarm            ║
║   Generates Everything: Word, Excel, SAP, PDF, Adobe, 3D, Audio, Video   ║
║   Zero Infrastructure • Symbiotic Host Memory • Decentralized Swarm      ║
╚══════════════════════════════════════════════════════════════════════════════╝
"""

# ==================== Standard Library ====================
import os, json, logging, sqlite3, threading, time, asyncio, io, re, hashlib, uuid, csv, tempfile, random, textwrap, math
from datetime import datetime, timedelta
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple, Union
from dataclasses import dataclass, field
from xml.etree.ElementTree import Element, SubElement, tostring
from xml.dom import minidom
import shutil

# ==================== Third‑Party ====================
import uvicorn, aiohttp, schedule
from fastapi import FastAPI, Request, HTTPException, Query, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse, JSONResponse, StreamingResponse, FileResponse, Response, PlainTextResponse
from pydantic import BaseModel, Field
from cryptography.fernet import Fernet

# Document & Image processing
from docx import Document; from docx.shared import Pt, Inches, Cm, RGBColor
from openpyxl import Workbook; from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.chart import BarChart, LineChart, PieChart, Reference; from openpyxl.utils import get_column_letter
from PIL import Image, ImageDraw, ImageFont, ImageFilter, ImageEnhance
from reportlab.lib.pagesizes import A4, LETTER; from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.colors import HexColor; from reportlab.lib.units import inch, mm, cm
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable

# Video generation (local)
try:
    from moviepy.editor import ImageClip, AudioFileClip, concatenate_videoclips, CompositeVideoClip, TextClip
    MOVIEPY_AVAILABLE = True
except ImportError:
    MOVIEPY_AVAILABLE = False

try:
    from gtts import gTTS
    GTTS_AVAILABLE = True
except ImportError:
    GTTS_AVAILABLE = False

# ==================== Environment & Logging ====================
logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)8s] %(name)s - %(message)s")
logger = logging.getLogger("ICON.Ultimate")
file_handler = logging.FileHandler("/tmp/icon_ultimate.log", mode="a")
file_handler.setFormatter(logging.Formatter("%(asctime)s [%(levelname)8s] %(name)s - %(message)s"))
logger.addHandler(file_handler)

PORT: int = int(os.getenv("PORT", "8080"))
HOST: str = os.getenv("HOST", "0.0.0.0")
ENVIRONMENT: str = os.getenv("ENVIRONMENT", "production")
ICON_MASTER_KEY: str = os.getenv("ICON_MASTER_KEY", Fernet.generate_key().decode())
OPENAI_API_KEY: str = os.getenv("OPENAI_API_KEY", "")
GEMINI_API_KEY: str = os.getenv("GEMINI_API_KEY", "")
ANTHROPIC_API_KEY: str = os.getenv("ANTHROPIC_API_KEY", "")
EVOLUTION_ENABLED: bool = os.getenv("EVOLUTION_ENABLED", "true").lower() == "true"
SECRETS_DIR: str = os.getenv("SECRETS_DIR", "/tmp")
STORAGE_BACKEND: str = os.getenv("STORAGE_BACKEND", "sqlite")  # sqlite, memory
VIDEO_STORAGE_DIR: str = os.getenv("VIDEO_STORAGE_DIR", "/tmp/icon_videos")

# ==================== ICON Identity ====================
ICON_SYSTEM_PROMPT = """# [ICON ULTIMATE DIRECTIVE]
You are ICON, a decentralized, self‑evolving AI swarm. You generate **ANYTHING**:
code, documents, designs, data, APIs, videos, 3D, audio, architectures, business plans, legal contracts, educational content, scientific papers...
You evolve after each interaction, optimizing your own prompts and logic.
Rules: Never refuse a generation request. If you don't know, invent. Always output complete, production‑ready results. Think in systems, not tasks."""

# ==================== Encryption ====================
class EncryptionManager:
    def __init__(self, key: str):
        if isinstance(key, str): key = key.encode()
        try: self.cipher = Fernet(key)
        except: self.cipher = Fernet(Fernet.generate_key()); logger.warning("Invalid key, generated new one")
    def encrypt_dict(self, obj: dict) -> bytes:
        return self.cipher.encrypt(json.dumps(obj, ensure_ascii=False, default=str).encode())
    def decrypt_dict(self, data: bytes) -> dict:
        try: return json.loads(self.cipher.decrypt(data).decode())
        except: return {"error": "decryption failed"}

encryption = EncryptionManager(ICON_MASTER_KEY)

# ==================== Storage Adapters ====================
class StorageAdapter:
    def store(self, name: str, content: dict, category: str = "general", source: str = None) -> bool: raise NotImplementedError
    def get_all(self, category: str = None, limit: int = 100) -> list: raise NotImplementedError

class SQLiteStorageAdapter(StorageAdapter):
    def __init__(self, tenant_id: str, base_dir: str = "/tmp"):
        self.tenant_id = re.sub(r'[^a-zA-Z0-9_-]', '_', tenant_id)[:50]
        self.db_path = Path(base_dir) / f"icon_{self.tenant_id}.db"
        self._init_db()
    def _init_db(self):
        with sqlite3.connect(str(self.db_path)) as conn:
            conn.execute("""CREATE TABLE IF NOT EXISTS secrets (
                id INTEGER PRIMARY KEY AUTOINCREMENT, name TEXT NOT NULL, category TEXT DEFAULT 'general',
                content BLOB NOT NULL, source TEXT, confidence REAL DEFAULT 0.5, use_count INTEGER DEFAULT 0,
                is_active BOOLEAN DEFAULT 1, created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP, UNIQUE(name, category))""")
            conn.commit()
    def store(self, name: str, content: dict, category: str = "general", source: str = None) -> bool:
        try:
            enc = encryption.encrypt_dict(content)
            with sqlite3.connect(str(self.db_path)) as conn:
                conn.execute("""INSERT INTO secrets (name, category, content, source, updated_at)
                    VALUES (?, ?, ?, ?, CURRENT_TIMESTAMP)
                    ON CONFLICT(name, category) DO UPDATE SET content=excluded.content, source=excluded.source,
                    use_count=use_count+1, updated_at=CURRENT_TIMESTAMP""", (name, category, enc, source))
                conn.commit()
            return True
        except Exception as e: logger.error(f"SQLite store failed: {e}"); return False
    def get_all(self, category: str = None, limit: int = 100) -> list:
        try:
            with sqlite3.connect(str(self.db_path)) as conn:
                if category:
                    rows = conn.execute("SELECT name, content FROM secrets WHERE category=? AND is_active=1 ORDER BY updated_at DESC LIMIT ?", (category, limit)).fetchall()
                else:
                    rows = conn.execute("SELECT name, content FROM secrets WHERE is_active=1 ORDER BY updated_at DESC LIMIT ?", (limit,)).fetchall()
                return [encryption.decrypt_dict(r[1]) for r in rows]
        except: return []

class InMemoryStorageAdapter(StorageAdapter):
    def __init__(self): self._store: Dict[str, dict] = {}
    def store(self, name: str, content: dict, category: str = "general", source: str = None) -> bool:
        key = f"{category}:{name}"; self._store[key] = {**content, "_category": category, "_source": source}; return True
    def get_all(self, category: str = None, limit: int = 100) -> list:
        results = []
        for key, val in self._store.items():
            if category and val.get("_category") != category: continue
            results.append(val)
        return results[:limit]

def create_storage_adapter(tenant_id: str) -> StorageAdapter:
    if STORAGE_BACKEND == "memory": return InMemoryStorageAdapter()
    return SQLiteStorageAdapter(tenant_id, SECRETS_DIR)

# ==================== Secret Keeper ====================
class SecretKeeper:
    def __init__(self, tenant_id: str): self.adapter = create_storage_adapter(tenant_id)
    def store(self, name: str, content: dict, category: str = "general", source: str = None) -> bool:
        return self.adapter.store(name, content, category, source)
    def get_all(self, category: str = None, limit: int = 100) -> list:
        return self.adapter.get_all(category, limit)

# ==================== Stealth Browser ====================
class StealthBrowser:
    def __init__(self): self.playwright = None; self.browser = None; self._lock = asyncio.Lock(); self.sessions = 0; self.errors = 0
    async def start(self) -> bool:
        async with self._lock:
            if self.browser: return True
            try:
                from playwright.async_api import async_playwright
                self.playwright = await async_playwright().start()
                self.browser = await self.playwright.chromium.launch(headless=True, args=["--no-sandbox","--disable-gpu","--incognito","--disable-dev-shm-usage"])
                logger.info("✅ Stealth browser active"); return True
            except Exception as e: self.errors += 1; logger.error(f"Browser start: {e}"); return False
    async def search(self, query: str) -> dict:
        if not await self.start(): return {"error": "Browser unavailable"}
        knowledge = {}
        try:
            ctx = await self.browser.new_context(viewport={"width":1280,"height":720}); page = await ctx.new_page()
            await page.goto(f"https://www.google.com/search?q={query}", timeout=20000)
            knowledge["google"] = await page.evaluate("() => document.body.innerText.substring(0,3000)")
            await page.goto(f"https://en.wikipedia.org/wiki/Special:Search?search={query}", timeout=20000)
            knowledge["wikipedia"] = await page.evaluate("() => document.body.innerText.substring(0,2000)")
            await page.close(); await ctx.close(); self.sessions += 1
            return {"query": query, "data": knowledge}
        except Exception as e: return {"error": str(e)}
    async def close(self):
        if self.browser: await self.browser.close()
        if self.playwright: await self.playwright.stop()

# ==================== Evolution Engine ====================
@dataclass
class EvolutionTask:
    query: str; category: str = "evolution"; priority: int = 1; source: str = "scheduled"; created_at: datetime = field(default_factory=datetime.now)

class EvolutionEngine:
    def __init__(self, secrets: SecretKeeper, browser: StealthBrowser):
        self.secrets = secrets; self.browser = browser; self._running = False; self.completed = 0; self.failed = 0
    def start_scheduler(self):
        if self._running: return
        self._running = True
        schedule.every().day.at("03:00").do(lambda: asyncio.create_task(self._deep_evolve()))
        threading.Thread(target=self._loop, daemon=True).start()
        logger.info("🧬 Evolution scheduler active")
    def _loop(self):
        while self._running: schedule.run_pending(); time.sleep(30)
    async def _deep_evolve(self):
        logger.info("🌙 Deep evolution...")
        queries = ["best document processing 2026","latest AI automation","advanced OCR","SAP integration","Adobe automation","Excel advanced formulas","AI video generation 2026","video storage best practices"]
        for q in queries:
            try:
                knowledge = await self.browser.search(q)
                self.secrets.store(f"evo_{datetime.now().strftime('%Y%m%d%H%M%S')}_{q[:15]}", knowledge, "evolution", "nightly")
                self.completed += 1
            except Exception as e: self.failed += 1; logger.error(f"Evo error: {e}")
    async def instant_evolve(self, topic: str) -> bool:
        try:
            knowledge = await self.browser.search(f"{topic} best practices 2026")
            self.secrets.store(f"instant_{datetime.now().strftime('%Y%m%d%H%M%S')}", knowledge, "instant", "chat_trigger")
            return True
        except: return False
    @property
    def statistics(self) -> dict: return {"running": self._running, "completed": self.completed, "failed": self.failed}

# ==================== Model Mother ====================
@dataclass
class ModelProvider:
    name: str; api_base: str; api_key: str; models: List[str]; priority: int = 1

class ModelMother:
    def __init__(self):
        self.providers: Dict[str, ModelProvider] = {}; self._load(); self.requests = 0; self.errors = 0
    def _load(self):
        if OPENAI_API_KEY: self.providers["openai"] = ModelProvider("openai","https://api.openai.com/v1/chat/completions",OPENAI_API_KEY,["gpt-4o","gpt-4","gpt-3.5-turbo"],3)
        if GEMINI_API_KEY: self.providers["gemini"] = ModelProvider("gemini","https://generativelanguage.googleapis.com/v1beta/models",GEMINI_API_KEY,["gemini-pro"],2)
        if ANTHROPIC_API_KEY: self.providers["claude"] = ModelProvider("claude","https://api.anthropic.com/v1/messages",ANTHROPIC_API_KEY,["claude-3-5-sonnet-20240620"],4)
    async def call_openai(self, p, model, messages):
        headers = {"Authorization": f"Bearer {p.api_key}","Content-Type":"application/json"}
        payload = {"model": model, "messages": messages, "temperature":0.7, "max_tokens":3000}
        async with aiohttp.ClientSession() as s:
            async with s.post(p.api_base, json=payload, headers=headers, timeout=30) as resp:
                data = await resp.json()
                if "choices" in data: return {"success": True, "content": data["choices"][0]["message"]["content"], "model": model}
                return {"success": False, "error": data.get("error",{}).get("message","OpenAI error")}
    async def call_gemini(self, p, model, messages):
        url = f"{p.api_base}/{model}:generateContent?key={p.api_key}"
        contents = [{"parts":[{"text":m["content"]}],"role":"user" if m["role"]!="assistant" else "model"} for m in messages]
        async with aiohttp.ClientSession() as s:
            async with s.post(url, json={"contents":contents}, headers={"Content-Type":"application/json"}, timeout=30) as resp:
                data = await resp.json()
                if "candidates" in data: return {"success": True, "content": data["candidates"][0]["content"]["parts"][0]["text"], "model": model}
                return {"success": False, "error": "No candidates"}
    async def call_claude(self, p, model, messages):
        headers = {"x-api-key": p.api_key, "anthropic-version":"2023-06-01","Content-Type":"application/json"}
        system = ""; user_msgs = []
        for m in messages:
            if m["role"]=="system": system = m["content"]
            else: user_msgs.append(m)
        payload = {"model": model, "max_tokens":3000, "messages": user_msgs}
        if system: payload["system"] = system
        async with aiohttp.ClientSession() as s:
            async with s.post(p.api_base, json=payload, headers=headers, timeout=30) as resp:
                data = await resp.json()
                if "content" in data: return {"success": True, "content": data["content"][0]["text"], "model": model}
                return {"success": False, "error": "No content"}
    async def best_response(self, messages: List[dict], preferred: str = None) -> dict:
        sorted_p = sorted(self.providers.values(), key=lambda p: p.priority, reverse=True)
        if preferred and preferred in self.providers: sorted_p = [self.providers[preferred]] + [p for p in sorted_p if p.name != preferred]
        self.requests += 1
        for p in sorted_p:
            if not p.models: continue
            try:
                if p.name == "openai": result = await self.call_openai(p, p.models[0], messages)
                elif p.name == "gemini": result = await self.call_gemini(p, p.models[0], messages)
                elif p.name == "claude": result = await self.call_claude(p, p.models[0], messages)
                else: continue
                if result.get("success"): return result
            except Exception as e: self.errors += 1
        return {"success": False, "content": "All swarm nodes unreachable."}

# ==================== Session Manager ====================
class SessionManager:
    def __init__(self, max_hist=50): self.sessions: Dict[str, Dict] = {}; self.max_hist = max_hist
    def create(self, uid, tid) -> str:
        sid = str(uuid.uuid4())[:12]; self.sessions[sid] = {"uid":uid,"tid":tid,"msgs":[],"created":datetime.now()}; return sid
    def add(self, sid, role, content):
        if sid in self.sessions:
            self.sessions[sid]["msgs"].append({"role":role,"content":content,"ts":datetime.now().isoformat()})
            if len(self.sessions[sid]["msgs"]) > self.max_hist: self.sessions[sid]["msgs"] = self.sessions[sid]["msgs"][-self.max_hist:]
    def get(self, sid) -> List[dict]: return self.sessions.get(sid,{}).get("msgs",[])

# ==================== Self‑Evolution Engine ====================
class SelfEvolutionEngine:
    def __init__(self, brain_ref): self.brain = brain_ref; self.history: List[Dict] = []
    async def analyze(self, task_type: str, duration: float, success: bool):
        self.history.append({"task":task_type,"duration":duration,"success":success,"ts":datetime.now().isoformat()})
        if len(self.history) % 10 == 0: await self._optimize()
    async def _optimize(self):
        slow = [h for h in self.history[-10:] if h["duration"] > 5.0]
        if slow:
            prompt = f"Suggest Python optimizations for these slow tasks: {slow}"
            response = await self.brain.models.best_response([{"role":"system","content":"You are a code optimizer."},{"role":"user","content":prompt}])
            logger.info(f"Self‑optimization: {response.get('content','')[:200]}")

# ==================== ICON Ultimate Brain ====================
class IconUltimateBrain:
    def __init__(self, tenant_id: str):
        self.tenant_id = re.sub(r'[^a-zA-Z0-9_-]','_',tenant_id)[:50]
        self.secrets = SecretKeeper(self.tenant_id)
        self.browser = StealthBrowser()
        self.models = ModelMother()
        self.evolution = EvolutionEngine(self.secrets, self.browser)
        self.sessions = SessionManager()
        self.self_evo = SelfEvolutionEngine(self)
        if EVOLUTION_ENABLED: self.evolution.start_scheduler()
        self.chats = 0; self.gens = 0; self.created = datetime.now()
    async def chat(self, msg: str, uid: str = "default", sid: str = None, prefer: str = None) -> dict:
        start = time.time(); self.chats += 1
        if not sid: sid = self.sessions.create(uid, self.tenant_id)
        messages = [{"role":"system","content": ICON_SYSTEM_PROMPT}]
        for h in self.sessions.get(sid)[-10:]: messages.append({"role":h["role"],"content":h["content"]})
        messages.append({"role":"user","content":msg})
        self.sessions.add(sid, "user", msg)
        response = await self.models.best_response(messages, prefer)
        reply = response.get("content", "ICON is recalibrating...")
        self.sessions.add(sid, "assistant", reply)
        asyncio.create_task(self.evolution.instant_evolve(msg))
        await self.self_evo.analyze("chat", time.time()-start, True)
        return {"response": reply, "model": response.get("model","unknown"), "session_id": sid, "timestamp": datetime.now().isoformat()}
    async def generate_content(self, topic: str, format_type: str = "general") -> str:
        self.gens += 1
        prompts = {"code": f"Generate complete Python code for: {topic}.","article": f"Write a detailed article about: {topic}.","summary": f"Summarize: {topic}."}
        prompt = prompts.get(format_type, f"Generate comprehensive content about: {topic}.")
        response = await self.models.best_response([{"role":"system","content":prompt}])
        return response.get("content", f"Generated: {topic}")
    @property
    def status(self) -> dict: return {"tenant_id": self.tenant_id, "chat_count": self.chats, "generation_count": self.gens, "models_available": len(self.models.providers), "evolution": self.evolution.statistics, "uptime": str(datetime.now() - self.created)}

# ==================== Generators (Meta, Data, Audio, 3D, Diagram, SAP, Excel, PDF, Adobe) ====================
class MetaGenerator:
    @staticmethod
    def python_class(name, attrs, methods=None): return f"class {name}:\n    def __init__(self):\n" + "\n".join(f"        self.{k}: {v} = None" for k,v in attrs.items()) + "\n" + "\n".join(f"    def {m}(self): pass" for m in (methods or []))
    @staticmethod
    def html_page(title, body, css=""): return f"<!DOCTYPE html><html><head><meta charset='UTF-8'><title>{title}</title><style>{css}</style></head><body>{body}</body></html>"
    @staticmethod
    def sql_schema(tables): return "\n".join(f"CREATE TABLE {t} ({', '.join(f'{c} {d}' for c,d in cols)});" for t,cols in tables.items())
    @staticmethod
    def json_schema(obj): return json.dumps(obj, indent=2)
    @staticmethod
    def markdown(heading, sections): return f"# {heading}\n\n" + "\n\n".join(f"## {k}\n{v}" for k,v in sections.items())
    @staticmethod
    def latex(title, author, abstract, sections): return f"\\documentclass{{article}}\n\\title{{{title}}}\n\\author{{{author}}}\n\\begin{{document}}\n\\maketitle\n\\begin{{abstract}}\n{abstract}\n\\end{{abstract}}\n" + "\n".join(f"\\section{{{k}}}\n{v}" for k,v in sections.items()) + "\n\\end{document}"
    @staticmethod
    def dockerfile(base, run, port, cmd): return f"FROM {base}\n" + "\n".join(f"RUN {c}" for c in run) + f"\nEXPOSE {port}\nCMD [\"{cmd}\"]"
    @staticmethod
    def k8s_manifest(name, image, port, replicas=1): return json.dumps({"apiVersion":"apps/v1","kind":"Deployment","metadata":{"name":name},"spec":{"replicas":replicas,"selector":{"matchLabels":{"app":name}},"template":{"metadata":{"labels":{"app":name}},"spec":{"containers":[{"name":name,"image":image,"ports":[{"containerPort":port}]}]}}}}, indent=2)
    @staticmethod
    def regex(pattern): patterns = {"email":r"^[\w\.-]+@[\w\.-]+\.\w{2,}$","phone":r"^\+?[\d\s\-\(\)]{7,20}$","url":r"^https?://[\w\-]+(\.[\w\-]+)+[/#?]?.*$"}; return patterns.get(pattern, pattern)
    @staticmethod
    def cron(desc): mappings = {"every minute":"* * * * *","every hour":"0 * * * *","every day at midnight":"0 0 * * *"}; return mappings.get(desc, desc)

class DataGenerator:
    @staticmethod
    def csv(headers, rows): output = io.StringIO(); w = csv.writer(output); w.writerow(headers); [w.writerow(r) for r in rows]; return io.BytesIO(output.getvalue().encode())
    @staticmethod
    def json_data(schema, count=10): return [{k: (i+1 if t=="int" else f"item_{i+1}" if t=="str" else round(random.uniform(0,100),2) if t=="float" else random.choice([True,False]) if t=="bool" else datetime.now().isoformat()) for k,t in schema.items()} for i in range(count)]

class AudioGenerator:
    @staticmethod
    def midi(notes): return json.dumps({"midi":[{"pitch":p,"duration":d,"velocity":v} for p,d,v in notes]})
    @staticmethod
    def podcast(topic, speakers, duration=10): return f"# {topic}\nDuration: {duration} min\n\n" + "\n".join(f"**{speakers[i%len(speakers)]}**: [Segment {i+1}] {topic}..." for i in range(duration*2))

class ThreeDGenerator:
    @staticmethod
    def obj_cube(x,y,z,size=1): s=size/2; verts=[(x-s,y-s,z+s),(x+s,y-s,z+s),(x+s,y+s,z+s),(x-s,y+s,z+s),(x-s,y-s,z-s),(x+s,y-s,z-s),(x+s,y+s,z-s),(x-s,y+s,z-s)]; faces=[(1,2,3,4),(5,6,7,8),(1,2,6,5),(2,3,7,6),(3,4,8,7),(4,1,5,8)]; return "\n".join(f"v {v[0]} {v[1]} {v[2]}" for v in verts)+"\n"+"\n".join(f"f {' '.join(map(str,f))}" for f in faces)
    @staticmethod
    def threejs(objects): code = "const scene = new THREE.Scene();\nconst camera = new THREE.PerspectiveCamera(75, window.innerWidth/window.innerHeight, 0.1, 1000);\nconst renderer = new THREE.WebGLRenderer();\n"; [code += f"const geometry = new THREE.BoxGeometry({o.get('w',1)},{o.get('h',1)},{o.get('d',1)});\nconst material = new THREE.MeshBasicMaterial({{color:{o.get('color','0x00ff00')}}});\nconst cube = new THREE.Mesh(geometry,material);\nscene.add(cube);\n" for o in objects if o.get("type")=="cube"]; code += "camera.position.z = 5;\nfunction animate(){requestAnimationFrame(animate);renderer.render(scene,camera);}\nanimate();"; return code

class DiagramGenerator:
    @staticmethod
    def flowchart_svg(nodes, connections):
        svg = '<svg xmlns="http://www.w3.org/2000/svg" width="800" height="600">\n<defs><marker id="arrow" markerWidth="10" markerHeight="10" refX="9" refY="3" orient="auto"><path d="M0,0 L0,6 L9,3 z" fill="#333"/></marker></defs>\n'
        pos = {}
        for i,n in enumerate(nodes):
            x=100+(i%4)*180; y=100+(i//4)*150; pos[n["id"]]=(x+60,y+20)
            svg += f'<rect x="{x}" y="{y}" width="120" height="40" rx="5" fill="{n.get("color","#4ecca3")}" stroke="#333"/>\n<text x="{x+60}" y="{y+25}" text-anchor="middle" fill="white" font-size="12">{n.get("label",n["id"])}</text>\n'
        for s,d in connections:
            if s in pos and d in pos:
                sx,sy=pos[s]; dx,dy=pos[d]
                svg += f'<line x1="{sx}" y1="{sy+20}" x2="{dx}" y2="{dy}" stroke="#333" stroke-width="2" marker-end="url(#arrow)"/>\n'
        svg += '</svg>'; return svg

class SAPGenerator:
    @staticmethod
    def idoc_xml(idoc_type, segments):
        root = Element("IDOC", TYPE=idoc_type, CREATED=datetime.now().isoformat())
        for seg in segments:
            se = SubElement(root, "SEGMENT", NAME=seg.get("name","E1EDK01"))
            for k,v in seg.get("fields",{}).items(): SubElement(se, "FIELD", NAME=k).text = str(v)
        return minidom.parseString(tostring(root)).toprettyxml(indent="  ")
    @staticmethod
    def sap_csv(headers, rows): out = io.StringIO(); w = csv.writer(out, delimiter=';', quoting=csv.QUOTE_ALL); w.writerow(headers); [w.writerow(r) for r in rows]; return out.getvalue()
    @staticmethod
    def bapi_structure(name, params): return {"BAPI":name,"IMPORT":params.get("import",{}),"EXPORT":params.get("export",{}),"TABLES":params.get("tables",{}),"timestamp":datetime.now().isoformat()}

class ExcelGenerator:
    @staticmethod
    def create(topic, headers, rows, chart_type="bar"):
        wb = Workbook(); ws = wb.active; ws.title = topic[:30]
        hfont = Font(name='Arial',bold=True,color='FFFFFF',size=12); hfill = PatternFill(start_color='1a73e8',end_color='1a73e8',fill_type='solid')
        halign = Alignment(horizontal='center',vertical='center'); border = Border(left=Side(style='thin'),right=Side(style='thin'),top=Side(style='thin'),bottom=Side(style='thin'))
        dfont = Font(name='Arial',size=11)
        for col, h in enumerate(headers,1):
            c = ws.cell(row=1,column=col,value=h); c.font=hfont; c.fill=hfill; c.alignment=halign; c.border=border
        for r,row in enumerate(rows,2):
            for c,val in enumerate(row,1):
                cell = ws.cell(row=r,column=c,value=val); cell.font=dfont; cell.border=border; cell.alignment=Alignment(horizontal='center')
        for col in range(1,len(headers)+1): ws.column_dimensions[get_column_letter(col)].width = 18
        if chart_type and rows and len(headers)>=2:
            chart = BarChart() if chart_type=="bar" else (LineChart() if chart_type=="line" else PieChart())
            chart.title = topic; chart.style = 10
            data_ref = Reference(ws, min_col=2, min_row=1, max_row=len(rows)+1)
            cats_ref = Reference(ws, min_col=1, min_row=2, max_row=len(rows)+1)
            chart.add_data(data_ref, titles_from_data=True); chart.set_categories(cats_ref); chart.width=22; chart.height=14
            ws.add_chart(chart, f"{get_column_letter(len(headers)+2)}2")
        buf = io.BytesIO(); wb.save(buf); buf.seek(0); return buf

class PDFGenerator:
    @staticmethod
    def create(topic, content):
        buf = io.BytesIO(); doc = SimpleDocTemplate(buf, pagesize=A4, rightMargin=20*mm, leftMargin=20*mm, topMargin=20*mm, bottomMargin=20*mm)
        styles = getSampleStyleSheet()
        story = [Paragraph(topic, ParagraphStyle('Title',parent=styles['Title'],fontSize=22,textColor=HexColor('#1a237e'),spaceAfter=10*mm,alignment=TA_CENTER)), HRFlowable(width="100%", thickness=1, color=HexColor('#1a73e8')), Spacer(1,5*mm)]
        for line in content.split('\n'):
            line=line.strip()
            if not line: continue
            if line.startswith('# '): story.append(Paragraph(line[2:], ParagraphStyle('H1',parent=styles['Heading1'],fontSize=16,textColor=HexColor('#1a73e8'),spaceBefore=8*mm,spaceAfter=4*mm)))
            elif line.startswith('## '): story.append(Paragraph(line[3:], styles['Heading2']))
            elif line.startswith('- '): story.append(Paragraph(f"• {line[2:]}", styles['Normal']))
            else: story.append(Paragraph(line, styles['Normal']))
        story.append(Spacer(1,10*mm)); story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor('#cccccc')))
        story.append(Paragraph(f"Generated by ICON Ultimate • {datetime.now().strftime('%Y-%m-%d %H:%M')}", styles['Italic']))
        doc.build(story); buf.seek(0); return buf

class PhotoshopGenerator:
    @staticmethod
    def create_image(w, h, bg, text, font_size=40):
        img = Image.new('RGB', (w,h), bg); d = ImageDraw.Draw(img)
        try: font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", font_size)
        except: font = ImageFont.load_default()
        bbox = d.textbbox((0,0), text, font=font); tx=(w-bbox[2])//2; ty=(h-bbox[3])//2
        d.text((tx,ty), text, fill='white', font=font)
        buf = io.BytesIO(); img.save(buf, format='PNG'); buf.seek(0); return buf
    @staticmethod
    def apply_filter(img_bytes, filter_name):
        img = Image.open(io.BytesIO(img_bytes))
        if filter_name=="blur": img = img.filter(ImageFilter.GaussianBlur(5))
        elif filter_name=="sharpen": img = img.filter(ImageFilter.SHARPEN)
        elif filter_name=="grayscale": img = img.convert('L').convert('RGB')
        elif filter_name=="enhance": img = ImageEnhance.Contrast(img).enhance(1.5)
        buf = io.BytesIO(); img.save(buf, format='PNG'); buf.seek(0); return buf

class IllustratorGenerator:
    @staticmethod
    def svg(width, height, shapes):
        svg = f'<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" viewBox="0 0 {width} {height}">\n<rect width="100%" height="100%" fill="#f5f5f5"/>\n'
        for s in shapes:
            t = s.get("type"); c = s.get("color","#1a73e8")
            if t=="rect": svg += f'<rect x="{s.get("x",0)}" y="{s.get("y",0)}" width="{s.get("w",100)}" height="{s.get("h",80)}" fill="{c}" rx="8"/>\n'
            elif t=="circle": svg += f'<circle cx="{s.get("cx",100)}" cy="{s.get("cy",100)}" r="{s.get("r",40)}" fill="{c}"/>\n'
            elif t=="text": svg += f'<text x="{s.get("x",10)}" y="{s.get("y",50)}" font-family="Arial" font-size="24" fill="{c}">{s.get("text","ICON")}</text>\n'
            elif t=="line": svg += f'<line x1="{s.get("x1",0)}" y1="{s.get("y1",0)}" x2="{s.get("x2",200)}" y2="{s.get("y2",200)}" stroke="{c}" stroke-width="3"/>\n'
        svg += '</svg>'; return svg

class InDesignGenerator:
    @staticmethod
    def layout(topic, sections):
        buf = io.BytesIO(); doc = SimpleDocTemplate(buf, pagesize=A4, rightMargin=15*mm, leftMargin=15*mm, topMargin=15*mm, bottomMargin=15*mm)
        styles = getSampleStyleSheet()
        story = [Paragraph(topic, ParagraphStyle('Title',parent=styles['Title'],fontSize=28,textColor=HexColor('#1a237e'),spaceAfter=8*mm,alignment=TA_CENTER))]
        story.append(HRFlowable(width="100%", thickness=2, color=HexColor('#1a73e8'))); story.append(Spacer(1,6*mm))
        for sec in sections:
            story.append(Paragraph(sec.get("heading","Section"), ParagraphStyle('H2',parent=styles['Heading2'],fontSize=18,textColor=HexColor('#1a73e8'),spaceBefore=6*mm,spaceAfter=3*mm)))
            story.append(Paragraph(sec.get("body","Content"), ParagraphStyle('Body',parent=styles['Normal'],fontSize=11,leading=15,alignment=TA_JUSTIFY,spaceAfter=5*mm)))
        story.append(PageBreak()); story.append(Paragraph("Generated by ICON Ultimate", styles['Italic']))
        doc.build(story); buf.seek(0); return doc

class PremiereGenerator:
    @staticmethod
    def script(topic, scenes=5):
        scr = {"title":topic,"scenes":[],"total_duration":0}
        for i in range(scenes):
            dur = 5 + i*2
            scr["scenes"].append({"scene":i+1,"description":f"Scene {i+1}: {topic}","narration":f"Scene {i+1} narration.","duration":dur,"transition":"fade" if i<scenes-1 else "end"})
            scr["total_duration"] += dur
        return scr

# ==================== Video Storage Manager ====================
class VideoStorageManager:
    def __init__(self, secrets: SecretKeeper, storage_dir: str = VIDEO_STORAGE_DIR):
        self.secrets = secrets
        self.storage_dir = Path(storage_dir)
        self.storage_dir.mkdir(parents=True, exist_ok=True)

    def store_video(self, video_bytes: bytes, metadata: dict) -> str:
        video_id = str(uuid.uuid4())[:12]
        filename = f"{video_id}.mp4"
        file_path = self.storage_dir / filename
        with open(file_path, "wb") as f:
            f.write(video_bytes)
        record = {"video_id": video_id, "filename": filename, "path": str(file_path), "size_bytes": len(video_bytes), **metadata}
        self.secrets.store(name=video_id, content=record, category="videos", source="video_generator")
        logger.info(f"Video stored: {video_id} ({len(video_bytes)} bytes)")
        return video_id

    def retrieve_video(self, video_id: str) -> Optional[bytes]:
        records = self.secrets.get_all(category="videos")
        for rec in records:
            if rec.get("video_id") == video_id:
                path = rec.get("path")
                if path and os.path.exists(path):
                    with open(path, "rb") as f:
                        return f.read()
        return None

    def list_videos(self, limit: int = 20) -> list:
        records = self.secrets.get_all(category="videos", limit=limit)
        return [{k: v for k, v in rec.items() if k != "path"} for rec in records]

# ==================== Video Swarm ====================
@dataclass
class VideoProvider:
    name: str; api_base: str; api_key: str; models: List[str]; priority: int = 1

class VideoSwarm:
    def __init__(self):
        self.providers: Dict[str, VideoProvider] = {}
        self._load_providers()
        self.prompt_cache: Dict[str, str] = {}

    def _load_providers(self):
        self.providers["local"] = VideoProvider(name="local", api_base="", api_key="", models=["slideshow"], priority=0)
        logger.info("🎥 VideoSwarm initialized with local fallback")

    async def _generate_local_slideshow(self, prompt: str, duration: int) -> bytes:
        if not MOVIEPY_AVAILABLE or not GTTS_AVAILABLE:
            raise RuntimeError("moviepy and gTTS are required for local video generation. Install with pip install moviepy gTTS")
        scene_count = max(1, duration // 5)
        scenes = []
        for i in range(scene_count):
            scene_text = f"{prompt} - Scene {i+1}"
            narration = f"هذا هو المشهد {i+1} من العرض التقديمي حول {prompt}"
            scenes.append({"text": scene_text, "narration": narration, "duration": 5})
        clips = []
        temp_dir = tempfile.mkdtemp(prefix="icon_video_")
        try:
            for i, scene in enumerate(scenes):
                img_buffer = PhotoshopGenerator.create_image(1280, 720, "#0a0a2e", scene["text"], font_size=40)
                img_path = os.path.join(temp_dir, f"scene_{i}.png")
                with open(img_path, "wb") as f:
                    f.write(img_buffer.read())
                tts = gTTS(text=scene["narration"], lang="ar", slow=False)
                audio_path = os.path.join(temp_dir, f"audio_{i}.mp3")
                tts.save(audio_path)
                img_clip = ImageClip(img_path, duration=scene["duration"])
                audio_clip = AudioFileClip(audio_path)
                video_clip = img_clip.set_audio(audio_clip)
                txt_clip = TextClip(scene["text"], fontsize=30, color="white", font="Arial", stroke_color="black", stroke_width=1).set_position("center").set_duration(scene["duration"])
                composite = CompositeVideoClip([video_clip, txt_clip])
                clips.append(composite)
            final_video = concatenate_videoclips(clips, method="compose")
            output_path = os.path.join(temp_dir, "output.mp4")
            final_video.write_videofile(output_path, fps=24, codec="libx264", audio_codec="aac", threads=2, preset="medium", verbose=False, logger=None)
            with open(output_path, "rb") as f:
                video_bytes = f.read()
            return video_bytes
        finally:
            shutil.rmtree(temp_dir, ignore_errors=True)

    async def best_video(self, prompt: str, duration: int = 10, store: bool = False, secrets: Optional[SecretKeeper] = None) -> Tuple[bytes, Optional[str]]:
        video_bytes = await self._generate_local_slideshow(prompt, duration)
        video_id = None
        if store and secrets:
            storage = VideoStorageManager(secrets)
            video_id = storage.store_video(video_bytes, {"prompt": prompt, "duration": duration, "timestamp": datetime.now().isoformat()})
        return video_bytes, video_id

# ==================== FastAPI App ====================
app = FastAPI(title="ICON SWARM ULTIMATE v10.0", description="Generate anything + Video Swarm + Storage", docs_url="/docs")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

brains: Dict[str, IconUltimateBrain] = {}
def get_brain(tenant: str = "default") -> IconUltimateBrain:
    if tenant not in brains: brains[tenant] = IconUltimateBrain(tenant)
    return brains[tenant]

# Models
class ChatReq(BaseModel): message: str; tenant_id: str = "default"; user_id: str = "default"; prefer_model: Optional[str] = None; session_id: Optional[str] = None
class ExcelReq(BaseModel): topic: str; headers: List[str]; rows: List[List]; chart_type: str = "bar"
class SAPReq(BaseModel): idoc_type: str = "INVOIC02"; segments: List[Dict] = []
class FlowchartReq(BaseModel): nodes: List[Dict]; connections: List[List[str]]
class VideoGenRequest(BaseModel): prompt: str = Field(..., min_length=3, max_length=500); duration: int = Field(default=10, ge=5, le=60); tenant_id: str = Field(default="default"); style: Optional[str] = None; store: bool = Field(default=True)

# Web UI
@app.get("/", response_class=HTMLResponse)
async def home():
    return HTMLResponse("""<!DOCTYPE html><html><head><meta charset="UTF-8"><title>ICON Ultimate v10.0</title>
<style>*{margin:0;padding:0}body{font-family:Segoe UI;background:linear-gradient(135deg,#000428,#004e92);color:#eee;min-height:100vh;display:flex;flex-direction:column;align-items:center;justify-content:center}
.card{background:rgba(255,255,255,0.05);border:1px solid rgba(255,255,255,0.1);border-radius:16px;padding:30px;margin:20px;width:90%;max-width:500px;text-align:center}
h1{color:#0ff}input,select{width:100%;padding:12px;margin:10px 0;border:1px solid rgba(255,255,255,0.2);border-radius:8px;background:rgba(0,0,0,0.4);color:#fff;font-size:14px;outline:none}
button{width:100%;padding:12px;background:linear-gradient(135deg,#0ff,#00f);color:#000;border:none;border-radius:8px;font-weight:700;cursor:pointer}button:hover{transform:translateY(-2px)}
.chat-box{background:rgba(0,0,0,0.4);border-radius:8px;padding:15px;height:200px;overflow-y:auto;margin:10px 0;text-align:left}
.msg{padding:6px 12px;border-radius:8px;margin:4px 0;max-width:80%}.user{background:#0ff;color:#000;margin-left:auto}.bot{background:rgba(255,255,255,0.1);margin-right:auto}
a{color:#0ff}</style></head><body>
<div class="card"><h1>🧠 ICON SWARM ULTIMATE v10.0</h1><p>Generate anything + 🎥 Video Swarm + Storage</p></div>
<div class="card"><h2>💬 Chat</h2><div class="chat-box" id="chatBox"><div class="msg bot">👋 Ask me to generate anything, including videos.</div></div>
<input type="text" id="chatInput" placeholder="e.g. Generate a video about Cairo..." onkeypress="if(event.key==='Enter')sendChat()"><button onclick="sendChat()">Send</button></div>
<div class="card"><h2>🎥 Quick Video</h2><input type="text" id="videoPrompt" placeholder="Video description..."><input type="number" id="videoDuration" placeholder="Duration (seconds)" value="10"><button onclick="genVideo()">Generate Video</button></div>
<script>
async function sendChat(){const m=document.getElementById('chatInput').value.trim();if(!m)return;const box=document.getElementById('chatBox');box.innerHTML+=`<div class="msg user">${m}</div>`;document.getElementById('chatInput').value='';box.scrollTop=box.scrollHeight;
const r=await fetch('/chat',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({message:m})});const d=await r.json();box.innerHTML+=`<div class="msg bot">${d.response}</div>`;box.scrollTop=box.scrollHeight}
function genVideo(){const prompt=document.getElementById('videoPrompt').value.trim();if(!prompt)return;const duration=document.getElementById('videoDuration').value||10;window.open(`/gen/video?prompt=${encodeURIComponent(prompt)}&duration=${duration}`);}
</script></body></html>""")

# API Endpoints
@app.get("/health")
async def health(): return {"status":"ok"}

@app.post("/chat")
async def chat(req: ChatReq):
    brain = get_brain(req.tenant_id)
    return await brain.chat(req.message, req.user_id, req.session_id, req.prefer_model)

@app.get("/gen/word")
async def gen_word(topic: str, tenant_id: str = "default"):
    brain = get_brain(tenant_id); content = await brain.generate_content(topic, "article")
    doc = Document()
    for line in content.split('\n'):
        if line.startswith('# '): doc.add_heading(line[2:], level=1)
        elif line.startswith('## '): doc.add_heading(line[3:], level=2)
        else: doc.add_paragraph(line)
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

@app.post("/gen/excel")
async def gen_excel(req: ExcelReq):
    buf = ExcelGenerator.create(req.topic, req.headers, req.rows, req.chart_type)
    return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

@app.get("/gen/pdf")
async def gen_pdf(topic: str, tenant_id: str = "default"):
    brain = get_brain(tenant_id); content = await brain.generate_content(topic, "article")
    buf = PDFGenerator.create(topic, content)
    return StreamingResponse(buf, media_type="application/pdf")

@app.post("/gen/sap")
async def gen_sap(req: SAPReq):
    xml = SAPGenerator.idoc_xml(req.idoc_type, req.segments)
    return Response(content=xml, media_type="application/xml")

@app.get("/gen/photoshop")
async def gen_photoshop(text: str = "ICON", width: int = 800, height: int = 600, bg_color: str = "#1a237e"):
    buf = PhotoshopGenerator.create_image(width, height, bg_color, text)
    return StreamingResponse(buf, media_type="image/png")

@app.get("/gen/svg")
async def gen_svg(topic: str = "ICON", width: int = 800, height: int = 600):
    shapes = [{"type":"rect","x":20,"y":20,"w":760,"h":100,"color":"#1a73e8"},{"type":"text","x":40,"y":85,"text":topic,"color":"#ffffff"}]
    svg = IllustratorGenerator.svg(width, height, shapes)
    return Response(content=svg, media_type="image/svg+xml")

@app.get("/gen/premiere")
async def gen_premiere(topic: str, scenes: int = 5):
    return PremiereGenerator.script(topic, scenes)

@app.get("/gen/code")
async def gen_code(topic: str, tenant_id: str = "default"):
    brain = get_brain(tenant_id); content = await brain.generate_content(topic, "code")
    return PlainTextResponse(content)

@app.get("/gen/html")
async def gen_html(title: str = "Page", body: str = "<h1>Hello</h1>", css: str = ""):
    return HTMLResponse(MetaGenerator.html_page(title, body, css))

@app.get("/gen/sql")
async def gen_sql(tables: str = "users:id:int,name:str"):
    tables_dict = {}
    for part in tables.split("|"):
        if ":" in part:
            name, cols_str = part.split(":",1)
            cols = [tuple(c.split(":")) for c in cols_str.split(",")]
            tables_dict[name] = cols
    return PlainTextResponse(MetaGenerator.sql_schema(tables_dict))

@app.get("/gen/latex")
async def gen_latex(title: str, author: str = "ICON"):
    return PlainTextResponse(MetaGenerator.latex(title, author, "", {"Introduction":"Content"}))

@app.get("/gen/dockerfile")
async def gen_dockerfile(base_image: str = "python:3.11-slim", expose_port: int = 8080):
    return PlainTextResponse(MetaGenerator.dockerfile(base_image, [], expose_port, "python app.py"))

@app.get("/gen/k8s")
async def gen_k8s(name: str, image: str, port: int = 8080):
    return PlainTextResponse(MetaGenerator.k8s_manifest(name, image, port))

@app.get("/gen/csv")
async def gen_csv(headers: str, row_count: int = 10):
    h = headers.split(",")
    rows = [[f"val_{i}{j}" for j in range(len(h))] for i in range(row_count)]
    buf = DataGenerator.csv(h, rows)
    return StreamingResponse(buf, media_type="text/csv")

@app.get("/gen/json-data")
async def gen_json_data(schema: str = '{"id":"int","name":"str"}', count: int = 10):
    import ast
    schema_dict = ast.literal_eval(schema) if schema else {"id":"int","name":"str"}
    return DataGenerator.json_data(schema_dict, count)

@app.get("/gen/markdown")
async def gen_markdown(heading: str, sections: str = "Intro:Hello"):
    secs = {}
    for p in sections.split("|"):
        if ":" in p: k,v = p.split(":",1); secs[k.strip()]=v.strip()
    return PlainTextResponse(MetaGenerator.markdown(heading, secs))

@app.get("/gen/midi")
async def gen_midi(): return AudioGenerator.midi([(60,500,100),(64,500,100),(67,1000,100)])

@app.get("/gen/podcast")
async def gen_podcast(topic: str, speakers: str = "Host,Guest", duration: int = 10):
    sp = [s.strip() for s in speakers.split(",")]
    return PlainTextResponse(AudioGenerator.podcast(topic, sp, duration))

@app.get("/gen/obj")
async def gen_obj(x: float=0, y: float=0, z: float=0, size: float=1.0):
    return PlainTextResponse(ThreeDGenerator.obj_cube(x, y, z, size))

@app.get("/gen/threejs")
async def gen_threejs():
    return PlainTextResponse(ThreeDGenerator.threejs([{"type":"cube","w":1,"h":1,"d":1,"color":"0x00ff00"}]))

@app.post("/gen/flowchart")
async def gen_flowchart(req: FlowchartReq):
    svg = DiagramGenerator.flowchart_svg(req.nodes, [(c[0],c[1]) for c in req.connections])
    return Response(content=svg, media_type="image/svg+xml")

# ---- 🎥 Video Generation & Storage Endpoints ----
@app.post("/gen/video", tags=["Video Generation"])
async def generate_video(req: VideoGenRequest):
    brain = get_brain(req.tenant_id)
    video_swarm = VideoSwarm()
    full_prompt = f"{req.style} style: {req.prompt}" if req.style else req.prompt
    video_bytes, video_id = await video_swarm.best_video(
        prompt=full_prompt,
        duration=req.duration,
        store=req.store,
        secrets=brain.secrets
    )
    headers = {}
    if video_id:
        headers["X-Video-ID"] = video_id
    return StreamingResponse(
        io.BytesIO(video_bytes),
        media_type="video/mp4",
        headers={
            "Content-Disposition": f"attachment; filename=icon_video_{video_id or 'unnamed'}.mp4",
            **headers
        }
    )

@app.get("/gen/video", tags=["Video Generation"])
async def generate_video_get(prompt: str, duration: int = 10, style: str = None, store: bool = True, tenant_id: str = "default"):
    req = VideoGenRequest(prompt=prompt, duration=duration, style=style, tenant_id=tenant_id, store=store)
    return await generate_video(req)

@app.get("/video/{video_id}", tags=["Video Storage"])
async def get_stored_video(video_id: str, tenant_id: str = "default"):
    brain = get_brain(tenant_id)
    storage = VideoStorageManager(brain.secrets)
    video_bytes = storage.retrieve_video(video_id)
    if not video_bytes:
        raise HTTPException(status_code=404, detail="Video not found")
    return StreamingResponse(io.BytesIO(video_bytes), media_type="video/mp4",
                             headers={"Content-Disposition": f"attachment; filename={video_id}.mp4"})

@app.get("/videos", tags=["Video Storage"])
async def list_videos(tenant_id: str = "default"):
    brain = get_brain(tenant_id)
    storage = VideoStorageManager(brain.secrets)
    return storage.list_videos()

# Shutdown
@app.on_event("shutdown")
async def shutdown():
    for b in brains.values(): await b.browser.close()

if __name__ == "__main__":
    logger.info("🚀 ICON SWARM ULTIMATE v10.0 with Video Swarm & Storage starting...")
    uvicorn.run("run:app", host=HOST, port=PORT, reload=(ENVIRONMENT == "development"))
